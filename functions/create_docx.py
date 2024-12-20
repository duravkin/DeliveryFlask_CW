from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from flask import current_app
from models import OrderReport, Suborder


def generate_order_report(order_report_id):
    with current_app.app_context():
        # Получение данных отчета
        order_report = OrderReport.query.get(order_report_id)
        if not order_report:
            print(f"Отчет с ID {order_report_id} не найден.")
            return

        order = order_report.order
        suborders = Suborder.query.filter_by(order_id=order.id).all()

        # Создание документа
        doc = Document()

        # Вводный текст
        doc.add_heading(f"Отчет о заказе №{order_report_id}", level=1)
        intro_paragraph = doc.add_paragraph()
        intro_paragraph.add_run(
            "Добро пожаловать в компанию 'Доставка+'. Мы предлагаем надежные услуги "
            "по транспортировке грузов. Данный отчет содержит информацию о заказе и "
            "сопутствующих товарах."
        ).font.size = Pt(12)
        intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        doc.add_paragraph()

        # Информация о заказе
        doc.add_heading("Информация о заказе", level=2)
        doc.add_paragraph(
            f"Дата заказа: {order.order_date.strftime('%d.%m.%Y')}")
        doc.add_paragraph(f"Клиент: {order.client.client_name}")
        doc.add_paragraph(f"Маршрут: {order.route.departure_point} → {
                          order.route.destination_point}")
        doc.add_paragraph(f"Расстояние: {order.route.distance} км")
        doc.add_paragraph(f"Стоимость перевозки: {order.route.cost:.2f} руб.")
        doc.add_paragraph(f"Ревизор: {order_report.accountant.full_name}")
        doc.add_paragraph(f"Доход: {order_report.revenue:.2f} руб.")
        doc.add_paragraph(f"Расходы: {order_report.expenses:.2f} руб.")
        doc.add_paragraph(f"Прибыль: {order_report.profit:.2f} руб.")
        doc.add_paragraph()

        # Таблица с данными о товарах
        doc.add_heading("Товары в заказе", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Название товара"
        hdr_cells[1].text = "Категория"
        hdr_cells[2].text = "Цена за единицу"
        hdr_cells[3].text = "Количество"

        # Заполнение таблицы
        for suborder in suborders:
            product = suborder.product
            row_cells = table.add_row().cells
            row_cells[0].text = product.product_name
            row_cells[1].text = product.category
            row_cells[2].text = f"{product.price_per_unit:.2f} руб."
            row_cells[3].text = str(suborder.quantity)

        # Сохранение документа

        filepath = ".temp_docx/"
        filename = f"order_report_{order_report_id}.docx"
        doc.save(filepath + filename)
        print(f"Отчет успешно создан: {filename}")

        return filepath + filename
